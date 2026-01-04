from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_columns(section, cols):
    sectPr = section._sectPr
    cols_xml = sectPr.xpath('./w:cols')[0]
    cols_xml.set(qn('w:num'), str(cols))
    cols_xml.set(qn('w:space'), '720')  # 0.5 inch spacing

def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(10)
        run.font.bold = True
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(2)
    else:
        run.font.size = Pt(9)
        run.font.bold = True
        h.paragraph_format.space_before = Pt(4)
        h.paragraph_format.space_after = Pt(2)

def add_para(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(8)  # Condensed to fit max content
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(1)
    return p

def add_math_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(8)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

document = Document()
section = document.sections[0]
section.page_height = Inches(11.69)
section.page_width = Inches(8.27)
section.left_margin = Inches(0.25)
section.right_margin = Inches(0.25)
section.top_margin = Inches(0.3)
section.bottom_margin = Inches(0.3)

title = document.add_heading('AI Final Exam: Ultimate Cheat Sheet', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

set_columns(section, 2)

# --- 1. SEQUENTIAL DECISION PROCESSES ---
add_heading(document, '1. Sequential Decision Processes (MDP)')
add_para(document, 'Concept: Framework for decision making in stochastic environments. The "Markov Property" means the future depends only on the current state, not history.', bold=False)
add_para(document, 'The 5-Tuple Definition <S, A, T, R, γ>:', bold=True)
add_para(document, '• S (States): Set of all possible situations (e.g., Grid coords).')
add_para(document, '• A (Actions): Set of choices (e.g., N, S, E, W).')
add_para(document, '• T (Transition): P(s\'|s,a). Probability of reaching s\' given action a.')
add_para(document, '• R (Reward): R(s). Immediate feedback value.')
add_para(document, '• γ (Gamma): Discount factor (0≤γ≤1). 0=Myopic (Greedy), 1=Far-sighted.')

# --- 2. HMM FILTERING ---
add_heading(document, '2. HMM Filtering (Recursive Estimation)')
add_para(document, 'Concept: Computing the current belief state P(Xt | e1:t) given all evidence to date. Cycles Predict -> Update -> Normalize.')
add_para(document, 'The Formula:', bold=True)
add_math_para(document, 'bel(St) = α * P(et | St) * Σ [ P(St | st-1) * bel(st-1) ]')
add_para(document, 'English: Today\'s belief = Normalize [ (Likelihood of Evidence) * (Sum of all incoming Transition probs from yesterday) ]')
add_para(document, 'Variables: α = 1/Sum (Normalizer). P(e|S) = Sensor Model. P(S\'|S) = Transition.')

add_heading(document, 'Full Task: Student (Study, Soc, Rest)', level=2)
add_para(document, 'Given: Prior bel(S0) = [0.5, 0.3, 0.2].')
add_para(document, 'Trans T: [0.6 0.3 0.1; 0.2 0.5 0.3; 0.3 0.2 0.5] (Rows=From).')
add_para(document, 'Sensor P(L|S) = [0.7, 0.2, 0.1]. Obs: "Library" (L).')

add_para(document, 'Step 1: Predict (Transition from S0)', bold=True)
add_para(document, 'Formula: bel_bar(S1) = bel(S0) * T (Sum columns)')
add_para(document, '• To Study: (0.5*0.6) + (0.3*0.2) + (0.2*0.3) = 0.30+0.06+0.06 = 0.42')
add_para(document, '• To Soc:   (0.5*0.3) + (0.3*0.5) + (0.2*0.2) = 0.15+0.15+0.04 = 0.34')
add_para(document, '• To Rest:  (0.5*0.1) + (0.3*0.3) + (0.2*0.5) = 0.05+0.09+0.10 = 0.24')
add_para(document, 'Prediction Vector: [0.42, 0.34, 0.24]')

add_para(document, 'Step 2: Update (Sensor Model L)', bold=True)
add_para(document, 'Formula: Unnormalized = Prediction * Sensor')
add_para(document, '• Study: 0.42 * 0.7 = 0.294')
add_para(document, '• Soc:   0.34 * 0.2 = 0.068')
add_para(document, '• Rest:  0.24 * 0.1 = 0.024')

add_para(document, 'Step 3: Normalize', bold=True)
add_para(document, 'Sum = 0.294 + 0.068 + 0.024 = 0.386. α = 1/0.386.')
add_para(document, '• P(Study) = 0.294 / 0.386 ≈ 0.762')
add_para(document, '• P(Soc)   = 0.068 / 0.386 ≈ 0.176')
add_para(document, '• P(Rest)  = 0.024 / 0.386 ≈ 0.062')
add_para(document, 'Final Answer: [0.762, 0.176, 0.062]')

# --- 3. VALUE ITERATION ---
add_heading(document, '3. MDP Planning: Value Iteration')
add_para(document, 'Concept: Calculates optimal U(s) by iterating Bellman equations. Converges to unique optimal values.')
add_para(document, 'The Formula (Bellman Update):', bold=True)
add_math_para(document, 'Ui+1(s) = R(s) + γ * max_a Σ [ P(s\'|s,a) * Ui(s\') ]')
add_para(document, 'English: My value = Immediate Reward + Discounted Value of the BEST future action.')

add_heading(document, 'Full Task: Risk vs Safety', level=2)
add_para(document, 'Map: A -> (Go: 0.8->B, 0.2->A). B -> (Exit: 1.0->Term).')
add_para(document, 'Rewards: R(A)=-1, R(B)=-2, R(T)=0. γ=0.9.')

add_para(document, 'Iteration 0 (Start):', bold=True)
add_para(document, 'U0(A) = 0, U0(B) = 0')

add_para(document, 'Iteration 1 (Compute U1 from U0):', bold=True)
add_para(document, '• State A (Action Go): R(A) + γ[0.8*U0(B) + 0.2*U0(A)]')
add_para(document, '  = -1 + 0.9[(0.8*0) + (0.2*0)] = -1.')
add_para(document, '• State B (Action Exit): R(B) + γ[1.0*U0(T)]')
add_para(document, '  = -2 + 0.9[0] = -2.')
add_para(document, 'Result: U1(A) = -1, U1(B) = -2.')

add_para(document, 'Iteration 2 (Compute U2 from U1):', bold=True)
add_para(document, '• State A (Go): -1 + 0.9[ 0.8*U1(B) + 0.2*U1(A) ]')
add_para(document, '  = -1 + 0.9[ (0.8*-2) + (0.2*-1) ]')
add_para(document, '  = -1 + 0.9[ -1.6 - 0.2 ] = -1 + 0.9(-1.8) = -1 - 1.62 = -2.62.')
add_para(document, '• State B (Exit): -2 + 0.9[0] = -2.')
add_para(document, 'Result: U2(A) = -2.62, U2(B) = -2.')

add_para(document, 'Iteration 3 (Compute U3 from U2):', bold=True)
add_para(document, '• State A (Go): -1 + 0.9[ 0.8*-2 + 0.2*-2.62 ]')
add_para(document, '  = -1 + 0.9[ -1.6 - 0.524 ] = -1 + 0.9(-2.124) = -2.91.')
add_para(document, 'Result: U3(A) = -2.91, U3(B) = -2.')

# --- 4. POLICY ITERATION ---
add_heading(document, '4. MDP Planning: Policy Iteration')
add_para(document, 'Concept: Two steps. 1) Evaluation: Solve linear equations for fixed policy. 2) Improvement: Check if max_a gives better score.')
add_para(document, 'Formula (Evaluation): U(s) = R(s) + γ Σ P(s\'|π) U(s\')')
add_para(document, 'Formula (Improvement): π(s) = argmax_a Σ P(s\'|a) U(s\')')

add_heading(document, 'Full Task: Policy Correction', level=2)
add_para(document, 'Policy π0: A=Go, B=Loop (Loop: 0.5->A, 0.5->B).')
add_para(document, 'Params: R(A)=-1, R(B)=-2, γ=0.9.')

add_para(document, 'Phase 1: Evaluation (Solve System)', bold=True)
add_para(document, 'Write Equations for π0:')
add_para(document, '1. U(A) = -1 + 0.9[0.8U(B) + 0.2U(A)]')
add_para(document, '2. U(B) = -2 + 0.9[0.5U(A) + 0.5U(B)]')
add_para(document, 'Simplify:')
add_para(document, '1. U(A) = -1 + 0.72U(B) + 0.18U(A) -> 0.82U(A) - 0.72U(B) = -1')
add_para(document, '2. U(B) = -2 + 0.45U(A) + 0.45U(B) -> -0.45U(A) + 0.55U(B) = -2')
add_para(document, 'Solve (Substitution):')
add_para(document, 'From Eq2: 0.55U(B) = -2 + 0.45U(A) -> U(B) ≈ -3.64 + 0.818U(A)')
add_para(document, 'Plug into Eq1: 0.82U(A) - 0.72(-3.64 + 0.818U(A)) = -1')
add_para(document, '...Algebra... -> U(A) ≈ -15.7. Then U(B) ≈ -16.4.')

add_para(document, 'Phase 2: Improvement (Lookahead)', bold=True)
add_para(document, '• Check A (Alt: Stay): R(A) + γU(A) = -1 + 0.9(-15.7) = -15.13.')
add_para(document, '  Compare: -15.13 (Stay) > -15.7 (Go). NEW: π(A) = Stay.')
add_para(document, '• Check B (Alt: Exit): R(B) + γU(T) = -2 + 0.9(0) = -2.0.')
add_para(document, '  Compare: -2.0 (Exit) > -16.4 (Loop). NEW: π(B) = Exit.')
add_para(document, 'Final Policy: {A: Stay, B: Exit}')

# --- 5. MONTE CARLO ---
add_heading(document, '5. Passive RL: Monte Carlo')
add_para(document, 'Concept: Model-free. Learns from complete episodes. High variance, zero bias. Must wait for end of episode.')
add_para(document, 'Formula: U(s) ← U(s) + α(Gt - U(s))')
add_para(document, 'Gt = Sum of discounted rewards from time t onwards.')

add_heading(document, 'Task: 2 Episodes. γ=0.9', level=2)
add_para(document, 'Start: U(A)=0, U(B)=0.')

add_para(document, 'Episode 1: A ->(-1)-> B ->(-2)-> Exit', bold=True)
add_para(document, '• Calc Return Gt (Work Backwards):')
add_para(document, '  G(B) = -2 + 0 = -2.')
add_para(document, '  G(A) = -1 + 0.9(-2) = -2.8.')
add_para(document, '• Update (First visit):')
add_para(document, '  U(A) = -2.8. U(B) = -2.')

add_para(document, 'Episode 2: A ->(-1)-> A ->(-1)-> B ->(-2)-> Exit', bold=True)
add_para(document, '• Calc Return Gt:')
add_para(document, '  G(B) = -2.')
add_para(document, '  G(A_2nd) = -1 + 0.9(-2) = -2.8.')
add_para(document, '  G(A_1st) = -1 + 0.9(-2.8) = -1 - 2.52 = -3.52.')
add_para(document, '• Update (Average):')
add_para(document, '  U(A) = Average(-2.8, -3.52) = -3.16.')
add_para(document, '  U(B) = Average(-2, -2) = -2.')

# --- 6. TD LEARNING ---
add_heading(document, '6. Passive RL: TD Learning')
add_para(document, 'Concept: Updates *online* after every step. Bootstraps (uses estimate of next state). Lower variance than MC.')
add_para(document, 'Formula: U(s) ← U(s) + α [ Target - U(s) ]')
add_para(document, 'Target = R(s) + γ * U(s\')')

add_heading(document, 'Task: 2 Episodes. α=0.5, γ=0.9', level=2)
add_para(document, 'Start: U(A)=0, U(B)=0.')

add_para(document, 'Episode 1: A -> B -> Exit', bold=True)
add_para(document, '• Step A->B (R=-1): Target = -1 + 0.9*U(B) = -1 + 0 = -1.')
add_para(document, '  U(A) <- 0 + 0.5(-1 - 0) = -0.5.')
add_para(document, '• Step B->Exit (R=-2): Target = -2 + 0 = -2.')
add_para(document, '  U(B) <- 0 + 0.5(-2 - 0) = -1.0.')
add_para(document, 'Current Vals: U(A)=-0.5, U(B)=-1.0.')

add_para(document, 'Episode 2: A -> A -> B -> Exit', bold=True)
add_para(document, '• Step A->A (R=-1): Target = -1 + 0.9*U(A) = -1 + 0.9(-0.5) = -1.45.')
add_para(document, '  U(A) <- -0.5 + 0.5(-1.45 - (-0.5)) = -0.975.')
add_para(document, '• Step A->B (R=-1): Target = -1 + 0.9*U(B) = -1 + 0.9(-1.0) = -1.9.')
add_para(document, '  U(A) <- -0.975 + 0.5(-1.9 - (-0.975)) = -1.4375.')
add_para(document, '• Step B->Exit (R=-2): Target = -2 + 0 = -2.')
add_para(document, '  U(B) <- -1.0 + 0.5(-2 - (-1.0)) = -1.5.')

# --- 7. Q-LEARNING ---
add_heading(document, '7. Active RL: Q-Learning')
add_para(document, 'Concept: Off-policy control. Learns Action-Values Q(s,a). Assumes optimal future behavior (Max Q) even if exploring.')
add_para(document, 'Formula: Q(s,a) ← Q + α [ R + γ * max_a\' Q(s\', a\') - Q ]')

add_heading(document, 'Task: 2 Episodes. α=0.5, γ=0.9', level=2)
add_para(document, 'Start: All Q(s,a) = 0. (A actions: Go, Stay. B actions: Loop, Exit).')

add_para(document, 'Episode 1: A (Go) -> B (Exit) -> Term', bold=True)
add_para(document, '• A->B (R=-1). Look at B: Q(Loop)=0, Q(Exit)=0. Max=0.')
add_para(document, '  Target = -1 + 0.9(0) = -1.')
add_para(document, '  Q(A,Go) <- 0 + 0.5(-1 - 0) = -0.5.')
add_para(document, '• B->Term (R=-2). Max Term=0.')
add_para(document, '  Target = -2 + 0 = -2.')
add_para(document, '  Q(B,Exit) <- 0 + 0.5(-2 - 0) = -1.0.')

add_para(document, 'Episode 2: A (Stay) -> A (Go) -> B (Exit) -> Term', bold=True)
add_para(document, '• A(Stay)->A (R=-1). Look at A: Q(Go)=-0.5, Q(Stay)=0. Max=0.')
add_para(document, '  Target = -1 + 0.9(0) = -1.')
add_para(document, '  Q(A,Stay) <- 0 + 0.5(-1 - 0) = -0.5.')
add_para(document, '• A(Go)->B (R=-1). Look at B: Q(Loop)=0, Q(Exit)=-1.0. Max=0.')
add_para(document, '  Target = -1 + 0.9(0) = -1.')
add_para(document, '  Q(A,Go) <- -0.5 + 0.5(-1 - (-0.5)) = -0.75.')
add_para(document, '• B(Exit)->Term (R=-2). Target = -2.')
add_para(document, '  Q(B,Exit) <- -1.0 + 0.5(-2 - (-1.0)) = -1.5.')

# --- 8. ADP ---
add_heading(document, '8. Model-Based RL: ADP')
add_para(document, 'Concept: Adaptive Dynamic Programming. 1) Count data to learn Model P. 2) Solve Bellman equations using learned P.')

add_heading(document, 'Task: 3 Iterations', level=2)
add_para(document, 'Data: Tried "Go" 3 times. Results: B, A, B.')
add_para(document, 'Step 1: Learn Model', bold=True)
add_para(document, '• P(B|Go) = 2/3 = 0.67.')
add_para(document, '• P(A|Go) = 1/3 = 0.33.')
add_para(document, 'Assumption: R(A)=-1. U(B)=-2 (Dead end).')

add_para(document, 'Step 2: Solve (Start U(A)=0)', bold=True)
add_para(document, '• Iteration 1:')
add_para(document, '  U(A) = -1 + 0.9 [ 0.33(0) + 0.67(-2) ]')
add_para(document, '       = -1 + 0.9 [ -1.34 ] = -2.206.')
add_para(document, '• Iteration 2 (Use U(A)=-2.206):')
add_para(document, '  U(A) = -1 + 0.9 [ 0.33(-2.206) + 0.67(-2) ]')
add_para(document, '       = -1 + 0.9 [ -0.728 - 1.34 ] = -1 + 0.9(-2.068) = -2.86.')
add_para(document, '• Iteration 3 (Use U(A)=-2.86):')
add_para(document, '  U(A) = -1 + 0.9 [ 0.33(-2.86) + 0.67(-2) ]')
add_para(document, '       = -1 + 0.9 [ -0.94 - 1.34 ] = -3.05.')

# --- 9. MIDTERM RECAP (Weeks 1-8) ---
add_heading(document, '9. Midterm Recap: CSP, Search, Bayes', level=1)

add_para(document, 'Search (W3-4):', bold=True)
add_para(document, '• BFS: Queue. Optimal (cost=1). Complete. O(b^d).')
add_para(document, '• DFS: Stack. Not optimal. O(bm) space.')
add_para(document, '• A*: f(n)=g(n)+h(n). Optimal if h(n) admissible (<= true cost).')

add_para(document, 'CSP (Constraint Satisfaction W5):', bold=True)
add_para(document, '• Components: Variables (X), Domains (D), Constraints (C).')
add_para(document, '• Backtracking: DFS with variable assignment.')
add_para(document, '• MRV (Heuristic): Pick variable with fewest legal values.')
add_para(document, '• Degree (Heuristic): Pick var involved in most constraints.')
add_para(document, '• LCV (Heuristic): Pick value that rules out fewest neighbors.')
add_para(document, '• Forward Checking: Delete values from neighbors incompatible with X.')
add_para(document, '• AC-3: Propagate consistency. Xi is consistent w/ Xj if for every val in Xi, there is allowed val in Xj.')

add_para(document, 'Adversarial (W6):', bold=True)
add_para(document, '• Minimax: Maximize own score, assume foe minimizes.')
add_para(document, '• Alpha-Beta Pruning: Stop searching if current path is worse than best option elsewhere. α=Best Max, β=Best Min.')
add_para(document, '• Expectiminimax: Weighted average for chance nodes.')

add_para(document, 'Bayes Nets (W7-8):', bold=True)
add_para(document, '• Bayes Rule: P(A|B) = P(B|A)P(A)/P(B).')
add_para(document, '• Cond. Indep: P(X,Y|Z) = P(X|Z)P(Y|Z).')
add_para(document, '• D-Separation (Active Trails):')
add_para(document, '  1. Chain A->B->C: Blocked if B known.')
add_para(document, '  2. Common Cause A<-B->C: Blocked if B known.')
add_para(document, '  3. V-Structure A->B<-C: ACTIVE if B (or descendent) known.')

document.save('AI_Final_Exam_Ultimate_Cheat_Sheet.docx')